
from docx import Document
import pandas as pd
from io import BytesIO
from docx.shared import Inches
from openpyxl import load_workbook
import os

def transformar_archivo(archivo_subido, plantilla_path):
    nombre = archivo_subido.name.lower()

    if nombre.endswith('.docx'):
        return procesar_docx(BytesIO(archivo_subido.read()), plantilla_path)
    elif nombre.endswith('.xlsx'):
        return procesar_excel(BytesIO(archivo_subido.read()), plantilla_path)
    else:
        raise ValueError("Formato de archivo no soportado. Usa .docx o .xlsx")

def procesar_docx(docx_file, plantilla_path):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    }

    def iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    doc_origen = Document(docx_file)
    doc_destino = Document(plantilla_path)

    indice_final = {
        "I. Flujo del Proceso": ["Flujo del Proceso"],
        "II. Descripción de Actividades": ["Descripción de actividades"],
        "III. Roles y Responsabilidades": ["Roles y Responsabilidades"],
        "IV. Criterio de Aceptación": ["Criterio de Aceptación"],
        "V. Indicadores de Desempeño": ["Indicadores de Desempeño", "Metas de Desempeño"],
        "VI. Riesgos": ["Riesgos"],
        "VII. Información Documentada Mantenida (Referencias)": ["Información Documentada de Referencia"],
        "VIII. Información Documentada Conservada (Registros)": ["Información documentada conservada"],
        "IX. Anexos": ["Anexos"],
        "X. Registro de Cambios": ["Registro de cambio", "Registro de cambios"]
    }

    contenido_clasificado = {titulo: [] for titulo in indice_final}
    seccion_actual = None

    for block in iter_block_items(doc_origen):
        if isinstance(block, Paragraph):
            texto = block.text.strip()
            encontrado = False
            for titulo_destino, keywords in indice_final.items():
                if any(texto.lower().startswith(kw.lower()) for kw in keywords):
                    seccion_actual = titulo_destino
                    contenido_clasificado[seccion_actual].append(("parrafo", block))
                    encontrado = True
                    break
            if not encontrado and seccion_actual:
                contenido_clasificado[seccion_actual].append(("parrafo", block))
        elif isinstance(block, Table) and seccion_actual:
            contenido_clasificado[seccion_actual].append(("tabla", block))

    for titulo_destino, bloques in contenido_clasificado.items():
        doc_destino.add_heading(titulo_destino, level=1)
        for tipo, contenido in bloques:
            if tipo == "parrafo":
                p_origen = contenido
                p = doc_destino.add_paragraph()
                for run in p_origen.runs:
                    r = p.add_run(run.text)
                    r.bold = run.bold
                    r.italic = run.italic
                    r.underline = run.underline
            elif tipo == "tabla":
                tabla = contenido
                t = doc_destino.add_table(rows=len(tabla.rows), cols=len(tabla.columns))
                t.style = 'Table Grid'
                for i, row in enumerate(tabla.rows):
                    for j, cell in enumerate(row.cells):
                        t.cell(i, j).text = cell.text.strip()

    buffer = BytesIO()
    doc_destino.save(buffer)
    buffer.seek(0)
    return buffer

def procesar_excel(excel_file, plantilla_path):
    secciones = {
        "I. Flujo del Proceso": "Flujo",
        "II. Descripción de Actividades": "Descripción Proceso",
        "III. Roles y Responsabilidades": "Descripción Proceso",
        "IV. Criterio de Aceptación": "Productos y Servicios",
        "V. Indicadores de Desempeño": "Indicadores",
        "VI. Riesgos": "Gestión Riesgos",
        "VII. Información Documentada Mantenida (Referencias)": "Información Documentada",
        "VIII. Información Documentada Conservada (Registros)": "Información Documentada",
        "X. Registro de Cambios": "Control de Cambios"
    }

    def agregar_tabla_desde_df(documento, df):
        tabla = documento.add_table(rows=1, cols=len(df.columns))
        tabla.style = 'Table Grid'
        hdr_cells = tabla.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        for _, row in df.iterrows():
            row_cells = tabla.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
        documento.add_paragraph("")

    def agregar_parrafo_con_viñeta(doc, texto):
        p = doc.add_paragraph(f"• {texto}")
        p.style = doc.styles['Normal']

    doc = Document(plantilla_path)
    xls = pd.ExcelFile(excel_file)
    wb = load_workbook(excel_file)

    for seccion, hoja in secciones.items():
        doc.add_heading(seccion, level=1)
        df = pd.read_excel(xls, sheet_name=hoja)

        if seccion == "II. Descripción de Actividades":
            col_e = df.iloc[:, 4].dropna()
            for val in col_e:
                if isinstance(val, str) and val.strip() not in ["Cliente", "ComoCuando"]:
                    agregar_parrafo_con_viñeta(doc, val.strip())

        elif seccion == "III. Roles y Responsabilidades":
            col_c = df.iloc[:, 2].dropna()
            unicos = sorted(set(val for val in col_c if isinstance(val, str) and val.strip() not in ["Etapa", "Quien"]))
            for val in unicos:
                agregar_parrafo_con_viñeta(doc, val.strip())

        elif seccion == "VII. Información Documentada Mantenida (Referencias)":
            tabla = df.iloc[:, 0:3].dropna(how="all")
            tabla.columns = [f"Columna {i+1}" for i in range(tabla.shape[1])]
            agregar_tabla_desde_df(doc, tabla)

        elif seccion == "VIII. Información Documentada Conservada (Registros)":
            tabla = df.iloc[:, 4:10].dropna(how="all")
            tabla.columns = [f"Columna {i+1}" for i in range(tabla.shape[1])]
            agregar_tabla_desde_df(doc, tabla)

        elif seccion == "I. Flujo del Proceso":
            hoja_flujo = wb[hoja]
            if hasattr(hoja_flujo, '_images') and hoja_flujo._images:
                img = hoja_flujo._images[0]
                image_stream = BytesIO(img._data())
                doc.add_picture(image_stream, width=Inches(6))
            elif not df.empty:
                agregar_tabla_desde_df(doc, df)

        else:
            if not df.empty:
                agregar_tabla_desde_df(doc, df)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
